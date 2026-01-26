"""
Document processor for handling PDFs and other academic documents.

This module provides functionality to:
- Extract text and structure from PDFs
- Extract metadata (title, authors, DOI, etc.)
- Chunk documents for embedding and retrieval
- Support multiple document formats
"""

import re
import hashlib
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, cast
from dataclasses import dataclass

# Try to import optional libraries
try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

try:
    from unstructured.partition.auto import partition

    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    partition = None

try:
    import pypdf

    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    pypdf = None

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None


@dataclass
class DocumentChunk:
    """Represents a chunk of a processed document."""

    text: str
    metadata: Dict[str, Any]
    chunk_id: str
    source: str


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""

    title: str
    authors: List[str]
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    source_path: str
    doc_id: str


class DocumentProcessor:
    """
    Document processor for academic papers and research documents.

    Supports PDF, DOCX, and plain text files with intelligent chunking
    and metadata extraction optimized for research content.
    """

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        min_chunk_size: int = 100,
        prefer_unstructured: bool = True,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target size of each chunk in words
            chunk_overlap: Number of words to overlap between chunks
            min_chunk_size: Minimum chunk size in words
            prefer_unstructured: Use unstructured library if available
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.prefer_unstructured = prefer_unstructured

        # Check available libraries
        self.has_pymupdf = PYMUPDF_AVAILABLE
        self.has_unstructured = UNSTRUCTURED_AVAILABLE
        self.has_pypdf = PYPDF_AVAILABLE
        self.has_docx = DOCX_AVAILABLE

    def process_document(self, file_path: str | Path) -> ProcessedDocument:
        """
        Process a document file and extract structured content.

        Args:
            file_path: Path to the document file

        Returns:
            ProcessedDocument with extracted content and metadata
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Generate document ID
        doc_id = self._generate_doc_id(file_path)

        # Extract content based on file type
        if file_path.suffix.lower() == ".pdf":
            content, metadata = self._extract_pdf_content(file_path)
        elif file_path.suffix.lower() in [".docx", ".doc"]:
            content, metadata = self._extract_docx_content(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            content, metadata = self._extract_text_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        content = self._normalize_text(content)

        if not content.strip():
            raise ValueError(f"No content extracted from document: {file_path}")

        # Enhance metadata
        metadata.update(
            {
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix.lower(),
                "processed_at": datetime.now(timezone.utc).isoformat(),
                "doc_id": doc_id,
            }
        )

        metadata.update(self._extract_content_metadata(content))

        # Extract title and authors if not already present
        if not metadata.get("title"):
            metadata["title"] = self._extract_title(content)
        if not metadata.get("authors"):
            metadata["authors"] = self._extract_authors(content)

        # Create chunks
        chunks = self._create_chunks(content, metadata)

        # Build lightweight validation summary
        metadata["validation"] = self._build_validation_summary(
            content=content, metadata=metadata, chunks=chunks
        )

        return ProcessedDocument(
            title=metadata.get("title", "Unknown Title"),
            authors=metadata.get("authors", []),
            content=content,
            chunks=chunks,
            metadata=metadata,
            source_path=str(file_path),
            doc_id=doc_id,
        )

    def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF file."""
        metadata = {}

        # Try unstructured first if available and preferred
        if self.has_unstructured and self.prefer_unstructured:
            try:
                return self._extract_pdf_unstructured(file_path)
            except Exception as e:
                print(f"Unstructured PDF processing failed: {e}")

        # Fall back to PyMuPDF
        if self.has_pymupdf:
            try:
                return self._extract_pdf_pymupdf(file_path)
            except Exception as e:
                print(f"PyMuPDF processing failed: {e}")

        # Final fallback to pypdf
        if self.has_pypdf:
            return self._extract_pdf_pypdf(file_path)

        raise ImportError("No PDF processing library available")

    def _extract_pdf_unstructured(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract PDF content using unstructured library."""
        elements = cast(Any, partition)(filename=str(file_path))

        # Separate content by type
        text_content = []
        metadata = cast(Dict[str, Any], {"elements": []})

        for element in elements:
            text_content.append(str(element))
            element_data = {
                "type": str(getattr(element, "category", "Unknown")),
                "text": str(element)[:200],  # Preview
            }
            metadata["elements"].append(element_data)

        content = "\n\n".join(text_content)

        # Try to extract title from first title element
        for element in elements:
            if hasattr(element, "category") and element.category == "Title":
                metadata["title"] = str(element)
                break

        return content, metadata

    def _extract_pdf_pymupdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract PDF content using PyMuPDF (fitz)."""
        doc: Any = cast(Any, fitz).open(str(file_path))

        # Extract basic metadata
        doc_metadata = doc.metadata if doc.metadata else {}
        metadata: Dict[str, Any] = {
            "page_count": len(doc),
            "title": doc_metadata.get("title", ""),
            "author": doc_metadata.get("author", ""),
            "subject": doc_metadata.get("subject", ""),
            "creator": doc_metadata.get("creator", ""),
            "producer": doc_metadata.get("producer", ""),
        }

        # Extract text with structure
        full_text = ""
        sections = []
        current_section = {"title": "Introduction", "content": "", "page": 0}
        first_page_lines = []

        try:
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = cast(Any, page.get_text("dict")).get("blocks", [])

                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            spans = cast(Any, line).get("spans", [])
                            text = "".join(
                                [cast(Any, span).get("text", "") for span in spans]
                            )

                            # Try to detect headers by font size
                            if spans:
                                font_size = max(
                                    cast(Any, span).get("size", 0) for span in spans
                                )
                                is_bold = any(
                                    (cast(Any, span).get("flags", 0) & 2**4) != 0
                                    for span in spans
                                )

                                if page_num == 0 and text.strip():
                                    first_page_lines.append(
                                        (text.strip(), font_size, is_bold)
                                    )

                                # Header detection: larger font, bold, short text
                                if (
                                    (font_size > 12 or is_bold)
                                    and len(text.strip()) < 100
                                    and text.strip()
                                ):
                                    if current_section["content"]:
                                        sections.append(current_section)
                                    current_section = {
                                        "title": text.strip(),
                                        "content": "",
                                        "page": page_num + 1,
                                    }
                                else:
                                    current_section["content"] += text + " "

                            full_text += text + "\n"
        except Exception as e:
            print(f"Error extracting structured text: {e}")
            # Fallback to simple text extraction
            text_pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text_pages.append(page.get_text())
            full_text = "\n".join(text_pages)

        if current_section["content"]:
            sections.append(current_section)

        metadata["sections"] = sections

        if not metadata.get("title") or not metadata.get("authors"):
            title, authors = self._extract_title_authors_from_lines(first_page_lines)
            if title and not metadata.get("title"):
                metadata["title"] = title
            if authors and not metadata.get("authors"):
                metadata["authors"] = authors

        doc.close()
        return full_text.strip(), metadata

    def _extract_pdf_pypdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract PDF content using pypdf (fallback method)."""
        with open(file_path, "rb") as file:
            pdf_reader = cast(Any, pypdf).PdfReader(file)

            metadata: Dict[str, Any] = {
                "page_count": len(pdf_reader.pages),
            }

            # Extract metadata if available
            if pdf_reader.metadata:
                metadata.update(
                    {
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                    }
                )

            text_content = []
            for page in pdf_reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as e:
                    print(f"Error extracting text from page: {e}")

            content = "\n\n".join(text_content)

            return content, metadata

    def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX file."""
        doc = cast(Any, DocxDocument)(str(file_path))

        metadata: Dict[str, Any] = {}

        if hasattr(doc, "core_properties"):
            core = doc.core_properties
            if core.title:
                metadata["title"] = core.title
            if core.author:
                metadata["author"] = core.author

        # Extract paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        # Try to extract title (first paragraph that looks like a title)
        for para in paragraphs:
            if len(para) < 100 and para:
                metadata["title"] = para
                break

        # Extract tables if any
        if doc.tables:
            metadata["has_tables"] = True
            metadata["table_count"] = len(doc.tables)

        return content, metadata

    def _extract_text_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        metadata: Dict[str, Any] = {
            "word_count": len(content.split()),
            "char_count": len(content),
        }

        # Try to extract title (first line that looks like a title)
        lines = content.split("\n")
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                title = line.lstrip("#").strip()
                if title:
                    metadata["title"] = title
                    break
            elif len(line) < 100:
                metadata["title"] = line
                break

        return content, metadata

    def _create_chunks(
        self, content: str, metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Create chunks from document content with overlap.

        Args:
            content: Full document content
            metadata: Document metadata

        Returns:
            List of DocumentChunk objects
        """
        # Split into words
        words = content.split()
        chunks = []

        step = max(1, self.chunk_size - self.chunk_overlap)

        # Create overlapping chunks
        for i in range(0, len(words), step):
            chunk_words = words[i : i + self.chunk_size]
            chunk_text = " ".join(chunk_words)

            # Skip very small chunks unless it's the only one
            if len(chunk_words) < self.min_chunk_size and len(chunks) > 0:
                continue

            # Create chunk metadata
            chunk_metadata = metadata.copy()
            chunk_metadata.update(
                {
                    "chunk_index": len(chunks),
                    "word_count": len(chunk_words),
                    "char_count": len(chunk_text),
                    "start_word": i,
                    "end_word": min(i + len(chunk_words), len(words)),
                }
            )

            # Generate chunk ID
            chunk_id = f"{metadata.get('doc_id', 'unknown')}_chunk_{len(chunks)}"

            chunk = DocumentChunk(
                text=chunk_text,
                metadata=chunk_metadata,
                chunk_id=chunk_id,
                source=metadata.get("file_path", "unknown"),
            )

            chunks.append(chunk)

        return chunks

    def _extract_title(self, content: str) -> str:
        """Extract title from document content."""
        lines = content.split("\n")[:10]  # Check first 10 lines

        for line in lines:
            line = line.strip()
            if line.startswith("#"):
                line = line.lstrip("#").strip()
            # Look for short lines that might be titles
            if 10 < len(line) < 150 and not line.startswith("http"):
                # Skip if it looks like an email, URL, or reference
                if not any(
                    x in line.lower()
                    for x in ["@", "http", "doi:", "abstract", "introduction"]
                ):
                    return line

        return "Untitled Document"

    def _extract_authors(self, content: str) -> List[str]:
        """Extract author names from document content."""
        authors = []

        # Look for patterns that suggest author names
        lines = content.split("\n")[:20]  # Check first 20 lines

        for i, line in enumerate(lines):
            line = line.strip()

            if not line:
                continue

            if "author" in line.lower() and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line:
                    line = next_line

            # Look for lines with author-like patterns
            # Skip common non-author lines
            skip_patterns = [
                "abstract",
                "introduction",
                "keywords",
                "university",
                "department",
                "@",
                "http",
            ]
            if any(pattern in line.lower() for pattern in skip_patterns):
                continue

            # Look for patterns like "Author Name", "Author Name1, Author Name2"
            if 5 < len(line) < 200:
                if line.lower().startswith("by "):
                    line = line[3:].strip()

                if self._looks_like_author_line(line):
                    potential_authors = re.split(r",\s*|\s+and\s+", line)
                    for author in potential_authors:
                        author = author.strip()
                        if self._is_name_candidate(author):
                            authors.append(author)
                elif (
                    1 < len(line.split()) <= 4
                    and line[0].isupper()
                    and self._is_name_candidate(line)
                ):
                    authors.append(line)

            # Stop after we find some authors or reach abstract
            if authors or "abstract" in line.lower():
                break

        return authors[:5]  # Limit to first 5 authors

    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID."""
        # Use file path and modification time
        stat = file_path.stat()
        unique_string = f"{file_path}_{stat.st_size}_{stat.st_mtime}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:16]

    def _normalize_text(self, content: str) -> str:
        """Normalize extracted text for downstream processing."""
        text = content.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_title_authors_from_lines(self, lines: List[Tuple[str, float, bool]]):
        """Extract title and authors from first-page lines."""
        if not lines:
            return None, []

        max_size = max(size for _, size, _ in lines)
        title = None
        title_index = None
        authors = []

        title_lines = []
        for i, (text, size, _) in enumerate(lines):
            cleaned = text.strip()
            if not cleaned:
                continue
            if any(x in cleaned.lower() for x in ["abstract", "introduction"]):
                continue
            if size >= max_size - 0.5 and len(cleaned) < 200:
                title_lines.append((i, cleaned))

        if title_lines:
            title_index = title_lines[0][0]
            combined = []
            last_index = title_index - 1
            for idx, text in title_lines:
                if idx == last_index + 1 and len(" ".join(combined + [text])) < 200:
                    combined.append(text)
                    last_index = idx
                elif not combined:
                    combined.append(text)
                    last_index = idx
                else:
                    break
            title = " ".join(combined).strip()

        if title and "  " in title:
            parts = re.split(r"\s{2,}", title)
            if len(parts) > 1 and self._looks_like_author_line(parts[-1]):
                candidates = re.split(r",\s*|\s+and\s+", parts[-1].strip())
                for author in candidates:
                    author = author.strip()
                    if self._is_name_candidate(author):
                        authors.append(author)
                if authors:
                    title = " ".join(parts[:-1]).strip()

        if title_index is None:
            return None, []

        for text, _, _ in lines[title_index + 1 :]:
            cleaned = text.strip()
            if not cleaned:
                break
            if any(
                x in cleaned.lower() for x in ["abstract", "introduction", "keywords"]
            ):
                break
            if len(cleaned) > 200:
                break

            line = cleaned
            if line.lower().startswith("by "):
                line = line[3:].strip()

            if not self._looks_like_author_line(line):
                break

            candidates = re.split(r",\s*|\s+and\s+", line)
            for author in candidates:
                author = author.strip()
                if self._is_name_candidate(author):
                    authors.append(author)
            if authors:
                break

        if not authors:
            for text, _, _ in lines:
                cleaned = text.strip()
                if not cleaned:
                    continue
                if self._looks_like_author_line(cleaned):
                    candidates = re.split(r",\s*|\s+and\s+", cleaned)
                    for author in candidates:
                        author = author.strip()
                        if self._is_name_candidate(author):
                            authors.append(author)
                    if authors:
                        break

        return title, authors

    def _looks_like_author_line(self, text: str) -> bool:
        """Heuristic to detect likely author lines."""
        cleaned = text.strip()
        if not cleaned:
            return False
        if cleaned.endswith("."):
            return False
        if cleaned.isupper() and " " not in cleaned:
            return False
        if ":" in cleaned:
            return False
        if any(ch.isdigit() for ch in cleaned):
            return False

        blocked_tokens = [
            "university",
            "department",
            "institute",
            "school",
            "college",
            "@",
            "http",
            "report",
            "copyright",
            "association",
            "programme",
            "program",
            "suite",
            "blvd",
            "street",
            "st.",
            "avenue",
            "road",
            "rd.",
            "isbn",
            "rights",
            "reserved",
            "published",
            "website",
            "www",
            "org",
        ]
        if any(token in cleaned.lower() for token in blocked_tokens):
            return False

        tokens = re.split(r"[ ,]+", cleaned)
        tokens = [t for t in tokens if t and t.lower() != "and"]
        if len(tokens) < 2:
            return False

        has_separator = "," in cleaned or " and " in cleaned.lower()
        if not has_separator and len(tokens) != 2:
            return False

        particles = {"de", "da", "del", "la", "van", "von", "der", "di"}
        lower_starts = sum(
            1 for t in tokens if t[0].islower() and t.lower() not in particles
        )
        if lower_starts > 0:
            return False

        uppercase_tokens = sum(1 for t in tokens if t.isupper())
        if uppercase_tokens >= len(tokens):
            return False

        return True

    def _is_name_candidate(self, text: str) -> bool:
        """Check if a token looks like a person name."""
        cleaned = text.strip()
        if not cleaned:
            return False
        if len(cleaned) < 3 or len(cleaned) > 60:
            return False
        if any(ch.isdigit() for ch in cleaned):
            return False
        if ":" in cleaned:
            return False
        if any(
            token in cleaned.lower()
            for token in [
                "university",
                "department",
                "institute",
                "school",
                "college",
                "@",
                "http",
                "report",
                "rights",
                "reserved",
                "published",
                "association",
                "programme",
                "program",
                "suite",
                "blvd",
                "isbn",
                "website",
                "www",
                "org",
            ]
        ):
            return False
        if cleaned.isupper():
            return False
        if not any(ch.isalpha() for ch in cleaned):
            return False
        if not cleaned[0].isupper():
            return False
        if len(cleaned.split()) < 2 and "." not in cleaned:
            return False
        return True

    def _extract_content_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata heuristics from content."""
        normalized = " ".join(content.split())
        word_count = len(normalized.split())

        metadata: Dict[str, Any] = {
            "word_count": word_count,
            "char_count": len(content),
            "has_abstract": "abstract" in normalized.lower()[:2000],
            "has_references": "references" in normalized.lower()[-5000:],
        }

        doi_match = re.search(
            r"\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", normalized, flags=re.IGNORECASE
        )
        if doi_match:
            metadata["doi"] = doi_match.group().rstrip(".,;)")
            metadata["doi_source"] = "content"

        return metadata

    def _build_validation_summary(
        self, content: str, metadata: Dict[str, Any], chunks: List[DocumentChunk]
    ) -> Dict[str, Any]:
        """Create a simple validation summary for extracted documents."""
        word_count = metadata.get("word_count") or len(content.split())

        avg_chunk_words = 0
        if chunks:
            avg_chunk_words = int(
                sum(c.metadata.get("word_count", 0) for c in chunks) / len(chunks)
            )

        return {
            "word_count": word_count,
            "chunk_count": len(chunks),
            "avg_chunk_words": avg_chunk_words,
            "has_title": bool(metadata.get("title")),
            "has_authors": bool(metadata.get("authors")),
            "has_doi": bool(metadata.get("doi")),
            "has_abstract": bool(metadata.get("has_abstract")),
            "has_references": bool(metadata.get("has_references")),
        }

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = [".txt", ".md"]

        if self.has_pymupdf or self.has_unstructured or self.has_pypdf:
            formats.append(".pdf")

        if self.has_docx:
            formats.append(".docx")

        return sorted(formats)
